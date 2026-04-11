"""
rag/document_processor.py
──────────────────────────
Handles loading documents of multiple formats and splitting them
into overlapping chunks suitable for embedding.

Supports:  .pdf  .docx  .txt

Interview talking point:
  "Chunking with overlap ensures that sentences spanning a chunk
   boundary are still retrieved — the overlap acts like a sliding window."
"""

import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from utils.config import config
from utils.logger import get_logger

log = get_logger("DocumentProcessor")


def load_document(file_path: str) -> List[Document]:
    """
    Load a single file and return a list of LangChain Document objects.
    Each Document holds page_content (text) and metadata (source, page).
    """
    path = Path(file_path)
    ext  = path.suffix.lower()
    log.info(f"Loading document: {path.name} ({ext})")

    if ext == ".pdf":
        return _load_pdf(file_path)
    elif ext == ".docx":
        return _load_docx(file_path)
    elif ext == ".txt":
        return _load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .pdf, .docx, or .txt")


def chunk_documents(documents: List[Document]) -> List[Document]:
    """
    Split documents into overlapping chunks.

    RecursiveCharacterTextSplitter tries to split on:
      '\n\n' → paragraph  →  '\n' → line  →  ' ' → word  →  ''
    This keeps semantically related text together.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        length_function=len,
    )
    chunks = splitter.split_documents(documents)
    log.info(f"Split {len(documents)} docs → {len(chunks)} chunks "
             f"(size={config.CHUNK_SIZE}, overlap={config.CHUNK_OVERLAP})")
    return chunks


def load_all_from_directory(directory: str) -> List[Document]:
    """
    Load and chunk ALL supported documents in a directory.
    Used for initial knowledge base seeding.
    """
    all_chunks: List[Document] = []
    dir_path = Path(directory)

    if not dir_path.exists():
        log.warning(f"Directory not found: {directory}")
        return []

    for file_path in dir_path.iterdir():
        if file_path.suffix.lower() in {".pdf", ".docx", ".txt"}:
            try:
                docs   = load_document(str(file_path))
                chunks = chunk_documents(docs)
                all_chunks.extend(chunks)
            except Exception as e:
                log.error(f"Failed to load {file_path.name}: {e}")

    log.info(f"Loaded {len(all_chunks)} total chunks from {directory}")
    return all_chunks


# ── Private loaders ────────────────────────────────────────────────────────────

def _load_pdf(file_path: str) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    documents = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(Document(
                page_content=text,
                metadata={"source": file_path, "page": i + 1},
            ))
    return documents


def _load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument

    doc   = DocxDocument(file_path)
    text  = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(
        page_content=text,
        metadata={"source": file_path, "page": 1},
    )]


def _load_txt(file_path: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(
        page_content=text,
        metadata={"source": file_path, "page": 1},
    )]
